from fastapi import FastAPI, UploadFile, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
from docx import Document
from openai import OpenAI
import json
import logging
import traceback
import shutil
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from tenacity import retry, stop_after_attempt, wait_exponential
from config_manager import ConfigManager
from datetime import datetime

# 初始化日志记录器
logger = logging.getLogger(__name__)

# 初始化配置管理器
config = ConfigManager()
config.ensure_directories()

app = FastAPI(title="文档AI助手", version="1.0.0")

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 创建static目录（如果不存在）
static_dir = Path("static")
static_dir.mkdir(exist_ok=True)

# 挂载静态文件目录
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def root():
    """根路由重定向到index.html"""
    return FileResponse("static/index.html")

# 默认的CoT提示词
DEFAULT_COT_PROMPT = """你是一个专业的实验报告助手。请仔细阅读以下实验报告模板，并完成实验准备、实验过程和实验总结等部分的内容填充。

请注意：
1. 保持原有的文档结构和格式不变
2. 重点填充以下部分（如果存在）：
   - 实验准备：详细列出所需的环境、工具、软件等
   - 实验过程：详细记录操作步骤、遇到的问题和解决方案
   - 实验总结：总结实验心得、成果、存在的问题和改进建议
3. 其他已填写的内容（如实验目��、实验要求等）保持不变
4. 生成的内容应该专业、具体、符合实验主题
5. 确保填写的内容与实验主题高度相关
6. 如果遇到表格，保持表格结构不变，仅填充内容

文档内容如下：
{document_content}

额外要求：
{additional_requirements}

请直接输出完整的文档内容，保持原有格式，重点填充实验准备、过程和总结部分。对于不需要修改的部分（如实验目的、要求等），请保持原样。"""

def extract_document_content(doc: Document) -> str:
    """提取文档中的所有内容，包括表格"""
    content = []
    
    # 首先获取所有段落内容（不包括表格中的段落）
    table_paragraphs = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_paragraphs.update(cell.paragraphs)
    
    # 添加非表格段落的内容
    for paragraph in doc.paragraphs:
        if paragraph not in table_paragraphs and paragraph.text.strip():
            content.append(paragraph.text)
    
    # 添加表格内容
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                row_content.append(cell.text.strip())
            if any(cell.strip() for cell in row_content):  # 只添加非空行
                table_content.append('\t'.join(row_content))
        if table_content:  # 只添加非空表格
            content.append('\n'.join(table_content))
    
    return '\n'.join(content)

def modify_document(doc_path: Path, ai_response: str) -> Path:
    """在原文档的基础上修改内容"""
    try:
        # 创建输出文件路径
        output_path = config.output_dir / f"modified_{doc_path.name}"
        
        # 首先复制原文档
        shutil.copy2(doc_path, output_path)
        
        # 打开复制的文档
        doc = Document(output_path)
        
        # 将AI响应按段落分割
        ai_paragraphs = [p.strip() for p in ai_response.split('\n') if p.strip()]
        ai_paragraph_index = 0
        
        # 获取所有有效段落（非空段落）
        valid_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        
        # 遍历文档中的段落
        for paragraph in valid_paragraphs:
            if ai_paragraph_index >= len(ai_paragraphs):
                break
                
            # 如果是表格的一部分，跳过
            if any(paragraph in cell.paragraphs for table in doc.tables for row in table.rows for cell in row.cells):
                continue
                
            # 清除原段落的所有运行
            paragraph._element.clear_content()
            # 添加新内容，保持原有格式
            paragraph.add_run(ai_paragraphs[ai_paragraph_index])
            ai_paragraph_index += 1
        
        # 处理表格
        for table in doc.tables:
            if ai_paragraph_index >= len(ai_paragraphs):
                break
                
            table_content = ai_paragraphs[ai_paragraph_index].split('\n')
            if len(table_content) > 1:  # 确认是表格内容
                for i, row in enumerate(table.rows):
                    if i < len(table_content):
                        cells = table_content[i].split('\t')
                        for j, cell in enumerate(row.cells):
                            if j < len(cells) and cells[j].strip():
                                # 清除单元格内容
                                for paragraph in cell.paragraphs:
                                    paragraph._element.clear_content()
                                # 添加新内容
                                cell.paragraphs[0].add_run(cells[j].strip())
                ai_paragraph_index += 1
        
        # 保存修改后的文档
        doc.save(output_path)
        return output_path
    
    except Exception as e:
        logger.error(f"修改文档时出错: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def save_ai_response(response: str, original_filename: str) -> Path:
    """保存AI的原始响应到文件"""
    try:
        # 创建responses目录（如果不存在）
        responses_dir = Path("responses")
        responses_dir.mkdir(exist_ok=True)
        
        # 创建响应文件名（使用原始文件名和时间戳）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        response_filename = f"{timestamp}_{Path(original_filename).stem}_response.txt"
        response_path = responses_dir / response_filename
        
        # 保存响应内容
        with open(response_path, "w", encoding="utf-8") as f:
            f.write(response)
        
        return response_path
    except Exception as e:
        logger.error(f"保存AI响应失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
async def call_openai_api(client: OpenAI, messages: list, model: str, temperature: float):
    """调用OpenAI API的重试包装器"""
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature
        )
        return response
    except Exception as e:
        logger.error(f"OpenAI API 调用失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise

@app.post("/upload")
async def upload_file(
    file: UploadFile,
    prompt: Optional[str] = Form(None),
    openai_api_key: Optional[str] = Form(None),
    openai_api_base: Optional[str] = Form(None),
    model: Optional[str] = Form(None)
):
    try:
        # 验证文件
        if not config.validate_file_extension(file.filename):
            raise HTTPException(status_code=400, detail="不支持的文件格式")
        
        content = await file.read()
        if not config.validate_file_size(len(content)):
            raise HTTPException(status_code=400, detail="文件大小超过限制")
        
        # 保存上传的文件
        file_path = config.upload_dir / file.filename
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        logger.info(f"文件 {file.filename} 上传成功")
        
        # 读取Word文档内容
        try:
            doc = Document(file_path)
            document_content = extract_document_content(doc)
            logger.info(f"成功读取文档，内容长度: {len(document_content)} 字符")
        except Exception as e:
            logger.error(f"读取文档失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(status_code=400, detail=str(e))
        
        # 配置OpenAI
        openai_config = config.get_openai_config()
        api_key = openai_api_key or openai_config.get('api_key')
        if not api_key:
            raise HTTPException(status_code=400, detail="未提供OpenAI API密钥")
        
        api_base = openai_api_base or openai_config.get('api_base')
        selected_model = model or openai_config.get('model', 'gpt-3.5-turbo-16k')
        temperature = openai_config.get('temperature', 0.7)
        
        logger.info(f"使用模型: {selected_model}")
        
        # 创建OpenAI客户端
        client = OpenAI(
            api_key=api_key,
            base_url=api_base
        )
        
        # 准备提示词
        final_prompt = DEFAULT_COT_PROMPT.format(
            document_content=document_content,
            additional_requirements=prompt if prompt else "请填充实验准备、实验过程和实验总结部分，保持专业性和具体性。"
        )
        
        try:
            # 调用OpenAI API
            response = await call_openai_api(
                client=client,
                messages=[
                    {"role": "system", "content": "你是一个专业的实验报告助手，擅长编写编译原理相关的实验报告。"},
                    {"role": "user", "content": final_prompt}
                ],
                model=selected_model,
                temperature=temperature
            )
            
            # 解析响应
            ai_response = response.choices[0].message.content
            
            # 保存原始响应
            response_path = save_ai_response(ai_response, file.filename)
            logger.info(f"AI响应已保存到: {response_path}")
            
            # 在日志中只显示响应的前100个字符
            preview = ai_response[:100] + "..." if len(ai_response) > 100 else ai_response
            logger.info(f"AI响应预览: {preview}")
            
            # 修改文档
            output_path = modify_document(file_path, ai_response)
            
            return {
                "message": "处理成功",
                "original_content": document_content,
                "ai_response": ai_response,
                "output_file": output_path.name,
                "response_file": response_path.name
            }
        
        except Exception as e:
            logger.error(f"AI处理失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"AI处理失败: {str(e)}")
            
    except Exception as e:
        logger.error(f"处理请求失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{filename}")
async def download_file(filename: str):
    try:
        file_path = config.output_dir / filename
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="文件不存在")
        return FileResponse(file_path)
    except Exception as e:
        logger.error(f"下载文件失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/config")
async def get_config():
    """获取当前配置"""
    try:
        return {
            "openai": {
                "api_base": config.get_openai_config().get('api_base'),
                "model": config.get_openai_config().get('model'),
                "temperature": config.get_openai_config().get('temperature')
            },
            "file": {
                "max_size_mb": config.get_file_config().get('max_size_mb'),
                "allowed_extensions": config.get_file_config().get('allowed_extensions')
            }
        }
    except Exception as e:
        logger.error(f"获取配置失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

# 添加新的路由来获取保存的响应
@app.get("/responses/{filename}")
async def get_response(filename: str):
    try:
        response_path = Path("responses") / filename
        if not response_path.exists():
            raise HTTPException(status_code=404, detail="响应文件不存在")
        return FileResponse(response_path)
    except Exception as e:
        logger.error(f"获取响应文件失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    server_config = config.get_server_config()
    uvicorn.run(
        "main:app",
        host=server_config.get('host', '0.0.0.0'),
        port=server_config.get('port', 8001),
        reload=server_config.get('reload', True)
    ) 